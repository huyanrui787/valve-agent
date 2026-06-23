"""投标文件 Word 成稿导出。"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..engines import Quotation, Verdict
from ..models.tender import ParsedTender

if TYPE_CHECKING:
    from ..agents.bid_agent import BidPackage


def _add_inline_bold(paragraph, text: str) -> None:
    """将含 **bold** 标记的文本写入段落，粗体部分设为 bold run。"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def export_bid_docx(
    package: BidPackage,
    path: str | Path,
    *,
    tender: ParsedTender | None = None,
    customer: str = "",
) -> Path:
    """将投标应答包导出为 Word 初稿。"""
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    title = tender.brief.title if tender and tender.brief.title else "阀门投标技术文件"
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"型号:{package.product_code} {package.product_name}\n").bold = True
    meta.add_run(f"生成日期:{date.today().isoformat()}\n")
    if customer:
        meta.add_run(f"客户:{customer}\n")
    if tender and tender.brief.bid_deadline:
        meta.add_run(f"投标截止:{tender.brief.bid_deadline.isoformat()}\n")

    _CN = "一二三四五六七八九十"

    def _h(n: int, title: str) -> None:
        doc.add_heading(f"{_CN[n-1]}、{title}", level=1)

    sec = 1  # 章节计数器

    if tender and tender.brief.key_points:
        _h(sec, "招标要点清单"); sec += 1
        for p in tender.brief.key_points:
            doc.add_paragraph(p, style="List Bullet")

    if tender and tender.brief.risks:
        _h(sec, "废标风险预警"); sec += 1
        rt = doc.add_table(rows=1, cols=3)
        rt.style = "Table Grid"
        hdr = rt.rows[0].cells
        hdr[0].text = "等级"
        hdr[1].text = "条款摘要"
        hdr[2].text = "说明"
        for risk in tender.brief.risks[:15]:
            row = rt.add_row().cells
            row[0].text = risk.level
            row[1].text = risk.clause[:80]
            row[2].text = risk.summary[:120]

    _h(sec, "技术偏离表"); sec += 1
    dt = package.deviation_table
    t = doc.add_table(rows=1, cols=6)
    t.style = "Table Grid"
    heads = ["序号", "参数", "招标要求", "产品能力", "判定", "建议/证据"]
    for i, htext in enumerate(heads):
        t.rows[0].cells[i].text = htext
    for item in dt.items:
        row = t.add_row().cells
        row[0].text = str(item.seq)
        row[1].text = item.param
        row[2].text = item.requirement
        row[3].text = item.product_capability
        row[4].text = item.verdict.value + (" ★" if item.is_critical else "")
        detail = item.suggestion or (item.evidence[0] if item.evidence else "")
        row[5].text = detail
        if item.verdict is Verdict.NEGATIVE:
            for run in row[4].paragraphs[0].runs:
                run.font.bold = True

    _h(sec, "废标风险体检"); sec += 1
    for item in package.compliance_report.items:
        p = doc.add_paragraph()
        p.add_run(f"[{item.level.value}] {item.name}: ").bold = True
        p.add_run(item.detail)

    if package.matched_qualifications:
        _h(sec, "资质匹配"); sec += 1
        for cat, q in package.matched_qualifications.items():
            if q:
                doc.add_paragraph(
                    f"{cat}: {q.name} 有效期至 {q.valid_until} (编号 {q.cert_no})",
                    style="List Bullet",
                )
            else:
                doc.add_paragraph(f"{cat}: 未匹配到有效资质", style="List Bullet")

    if package.matched_records:
        _h(sec, "业绩匹配"); sec += 1
        for r in package.matched_records[:10]:
            doc.add_paragraph(
                f"{r.contract_date} {r.project_name} — {r.customer} "
                f"({r.valve_type}, 约 {r.amount/1e4:.0f} 万元)",
                style="List Bullet",
            )

    if package.tech_proposal:
        _h(sec, "技术方案概述"); sec += 1
        for para in package.tech_proposal.split("\n"):
            stripped = para.strip()
            if not stripped:
                continue
            # Markdown 标题 → Word heading
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=2)
            # 带序号的大标题：一、 二、 三、...
            elif len(stripped) > 2 and stripped[1] == "、":
                h = doc.add_heading(stripped, level=2)
            # Markdown 列表项（– / - / • / ✅ / ⚠️ 开头）
            elif stripped.startswith(("– ", "- ", "• ", "✅ ", "⚠️ ", "* ")):
                bullet_text = stripped.lstrip("–-•✅⚠️* ").strip()
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_bold(p, bullet_text)
            else:
                p = doc.add_paragraph()
                _add_inline_bold(p, stripped)

    doc.add_paragraph()
    foot = doc.add_paragraph("— 本文件由 valve-agent 自动生成,供工程师复核定稿 —")
    foot.runs[0].font.size = Pt(9)
    doc.save(str(out))
    return out


def export_quotation_docx(quote: Quotation, path: str | Path) -> Path:
    """将整单报价导出为 Word 报价单。"""
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("阀门产品报价单", level=0)
    doc.add_paragraph(f"客户:{quote.customer or '(未填)'}")
    doc.add_paragraph(f"客户等级:{quote.customer_tier}  报价基准日:{quote.price_basis}")
    doc.add_paragraph(f"出口报价:{'是' if quote.is_export else '否'}")

    t = doc.add_table(rows=1, cols=7)
    t.style = "Table Grid"
    cols = ["型号", "名称", "DN", "数量", "含税单价", "小计", "毛利率"]
    for i, c in enumerate(cols):
        t.rows[0].cells[i].text = c
    for line in quote.lines:
        row = t.add_row().cells
        row[0].text = line.product_code
        row[1].text = line.product_name
        row[2].text = str(line.dn)
        row[3].text = str(line.quantity)
        row[4].text = f"{line.unit_price:,.2f}"
        row[5].text = f"{line.line_total:,.2f}"
        row[6].text = f"{line.margin:.1%}"

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"整单合计: ¥{quote.total:,.2f}  ").bold = True
    p.add_run(f"综合毛利: {quote.overall_margin:.1%}")

    doc.save(str(out))
    return out
