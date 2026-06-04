"""招标文件文本加载(PDF / Word / 纯文本)。"""

from __future__ import annotations

from pathlib import Path

_SUFFIXES = {".pdf", ".docx", ".doc", ".txt", ".md"}


def load_document(path: str | Path) -> str:
    """从文件读取纯文本。支持 .pdf / .docx / .txt / .md。"""
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"文件不存在:{p}")
    suffix = p.suffix.lower()
    if suffix not in _SUFFIXES:
        raise ValueError(f"不支持的文件类型:{suffix},支持 {_SUFFIXES}")

    if suffix == ".pdf":
        return _load_pdf(p)
    if suffix in (".docx", ".doc"):
        return _load_docx(p)
    return p.read_text(encoding="utf-8")


def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())
    return "\n\n".join(parts)


def _load_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
